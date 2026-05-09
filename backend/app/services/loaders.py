from typing import List
from langchain_core.documents import Document
import docx
import openpyxl
import os
import sys

# Try to import win32com for fallback
try:
    import win32com.client
    import pythoncom
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False

class CustomDocxLoader:
    def __init__(self, file_path: str):
        self.file_path = file_path
    
    def load(self) -> List[Document]:
        if not os.path.exists(self.file_path):
            raise RuntimeError(f"File not found: {self.file_path}")
            
        try:
            # Try standard python-docx first
            doc = docx.Document(self.file_path)
            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            # Add tables text as well
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text.append(" | ".join(row_text))
                    
            return [Document(page_content="\n".join(text), metadata={"source": self.file_path})]
        except Exception as e:
            # Check if it might be an encryption/package issue and we have Win32
            print(f"python-docx failed: {e}. Trying Win32 COM fallback...")
            
            if HAS_WIN32:
                try:
                    return self._load_with_win32()
                except Exception as w32_e:
                    raise RuntimeError(f"Error loading DOCX {self.file_path}: python-docx error: {e}; Win32 error: {w32_e}")
            else:
                raise RuntimeError(f"Error loading DOCX {self.file_path}: {e} (Win32 COM not available)")

    def _load_with_win32(self) -> List[Document]:
        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            
            # Use absolute path for COM
            abs_path = os.path.abspath(self.file_path)
            doc = word.Documents.Open(abs_path, ReadOnly=True)
            
            # Extract text
            text = doc.Range().Text
            
            # Clean up text (Word uses \r)
            text = text.replace('\r', '\n')
            
            return [Document(page_content=text, metadata={"source": self.file_path})]
        finally:
            if doc:
                try:
                    doc.Close(SaveChanges=False)
                except:
                    pass
            if word:
                try:
                    word.Quit()
                except:
                    pass
            pythoncom.CoUninitialize()

class CustomDocLoader:
    """加载 Word 97-2003 (.doc) 格式文件，使用 pypandoc 或 Windows COM。"""
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        if not os.path.exists(self.file_path):
            raise RuntimeError(f"File not found: {self.file_path}")
        abs_path = os.path.abspath(self.file_path)

        # 1. 尝试 pypandoc
        try:
            import pypandoc
            text = pypandoc.convert_file(abs_path, "plain", extra_args=["--wrap=none"])
            text = text.strip() if text else ""
            if text:
                return [Document(page_content=text, metadata={"source": self.file_path})]
        except ImportError:
            pass
        except Exception as e:
            print(f"pypandoc failed for .doc: {e}. Trying Win32 COM...")

        # 2. Windows 回退：使用 Win32 COM
        if HAS_WIN32:
            return self._load_doc_with_win32()
        raise RuntimeError(
            "无法解析 .doc 文件。请安装 pypandoc 和 Pandoc (pip install pypandoc)，"
            "或在使用 Windows 时安装 pywin32。"
        )

    def _load_doc_with_win32(self) -> List[Document]:
        pythoncom.CoInitialize()
        word = None
        doc_obj = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            abs_path = os.path.abspath(self.file_path)
            doc_obj = word.Documents.Open(abs_path, ReadOnly=True)
            text = doc_obj.Range().Text.replace("\r", "\n").strip()
            return [Document(page_content=text, metadata={"source": self.file_path})]
        finally:
            if doc_obj:
                try:
                    doc_obj.Close(SaveChanges=False)
                except Exception:
                    pass
            if word:
                try:
                    word.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()


class CustomExcelLoader:
    def __init__(self, file_path: str):
        self.file_path = file_path
        
    def load(self) -> List[Document]:
        try:
            wb = openpyxl.load_workbook(self.file_path, data_only=True)
            text = []
            for sheet in wb.worksheets:
                text.append(f"--- Sheet: {sheet.title} ---")
                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(cell) for cell in row if cell is not None]
                    if row_text:
                        text.append(" ".join(row_text))
            return [Document(page_content="\n".join(text), metadata={"source": self.file_path})]
        except Exception as e:
             raise RuntimeError(f"Error loading Excel {self.file_path}: {e}")
